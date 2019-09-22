import sys

try:
    import xml.etree.cElementTree as ET
except ImportError:
    import xml.etree.ElementTree as ET

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.text import WD_COLOR_INDEX

from docx.enum.dml import MSO_THEME_COLOR
###

#①②③④⑤⑥⑦⑧
DICT_ACCENT = {'0':'⓪',	'1':'①',	'2':'②',	'3':'③',	'4':'④',	'5':'⑤',	'6':'⑥',	'7':'⑦',	'8':'⑧'}
DICT_POS= {'N':'名詞',	'P':'代詞',	'NA':'形容動詞',	'adj':'adj',	'no':'~の',	'adv':'adv',
				'j5':'自五',	'j1':'自一',	't5':'他五',	't1':'他一',	'js':'自サ',	'ts':'他サ',	'jts':'自他サ',
				'5i':'自五',	'1i':'自一',	'5t':'他五',	'1t':'他一',
				'jt1':'自他一', 'jt5':'自他五'}

def dump_accent(paragraph, ac):
	if ac:
		error_prone = False
		str = ''
		for a in ac:
			try:
				str += DICT_ACCENT[a]
			except:
				error_prone=True
		if error_prone:
			paragraph.add_run(str).font.color.rgb = RGBColor(0xf2, 0x24, 0x24)
		else:
			paragraph.add_run(str)
	return

def dump_ex(doc, ex):
	word = ex.attrib['word']
	print(word)
	pron = ex.attrib['pron']
	pa = ex.get('pa')
	def dump_pos(pos):
		str = ''
		l = len(pos.split('/'))
		i = 0
		for p in pos.split('/'):
			str += DICT_POS[p]
			i+=1
			if i != l:
				str += '・'
		return str

	paragragh = doc.add_paragraph()
	paragragh.add_run(word+'（' + pron)
	dump_accent(paragragh, pa)
	paragragh.add_run('）')
	paragragh.paragraph_format.left_indent = Inches(0.8)
	jlpt = ex.get('jlpt')
	if jlpt:
		jlpt = 'N'+jlpt
		paragragh.add_run(jlpt).font.highlight_color =  WD_COLOR_INDEX.TURQUOISE
	pos = ex.get('pos')
	if pos:
		paragragh.add_run(dump_pos(pos)).font.highlight_color =  WD_COLOR_INDEX.YELLOW
	return

def dump_r(doc, r):
	on = r.get('on')
	kun = r.get('kun')
	jyg = r.get('jy')=='false'
	
	def fix_uncommon(p):
		for run in p.runs:
			run.font.strike = True
		return
	
	if on:
		paragragh = doc.add_paragraph(on)
		paragragh.paragraph_format.left_indent = Inches(0.5)
		paragragh.runs[0].font.color.rgb = RGBColor(0xf2, 0x24, 0x24)
		for child in r:
			dump_ex(doc, child)
	if kun:
		if 0 <= kun.find('`'):
			ss = kun.split('`')
			paragragh = doc.add_paragraph()
			paragragh.add_run(ss[0]).font.color.rgb = RGBColor(0xf2, 0x24, 0x24)
			paragragh.add_run(ss[1])
		else:
			paragragh = doc.add_paragraph(kun)
			paragragh.runs[0].font.color.rgb = RGBColor(0xf2, 0x24, 0x24)
			
		if jyg:
			fix_uncommon(paragragh)	
		paragragh.paragraph_format.left_indent = Inches(0.5)
		kk = r.get('kk')
		pa = r.get('pa')
		if pa:
			dump_accent(paragragh, pa)
		if kk:
			run = paragragh.add_run(' '+kk)

		for child in r:
			dump_ex(doc, child)
	return 0

def dump_kanji(doc, kanji):
	paragragh = doc.add_paragraph()
	run = paragragh.add_run(kanji.attrib['id'])
	print(kanji.attrib['id'])
	run.font.size = Pt(36)
	freq = kanji.attrib['freq']
	fn = int(freq)
	if fn < 1000:
		freq = '0'+freq
	if fn < 100:
		freq = '0'+freq
	run = paragragh.add_run(freq).font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
	jlpt = kanji.get('jlpt')
	if jlpt:
		jlpt = 'N'+jlpt
	else:
		jlpt = 'N-'
	gr= kanji.get('gr')
	if gr:
		gr = gr + '年'
	else:
		gr = '-年'
	run = paragragh.add_run(jlpt+'	'+gr)
	
	run.font.highlight_color =  WD_COLOR_INDEX.TURQUOISE
	
	for r in kanji:
		dump_r(doc, r)
	doc.add_paragraph()

count = 0

def dump_all(xml_name, doc_name):
	doc = Document()
	f = open(xml_name, 'rb')
	file_context=f.read() 
	f.close()
	
	tree = ET.ElementTree(file=xml_name)
	root = tree.getroot()
	print(root.tag, root.attrib)
	count = 0
	for child in root:
		count = count + 1
		dump_kanji(doc, child)
	
	doc.save(doc_name)
	
#dump_all('kanji.xml', "常用漢字.docx")
"""
dump_all('kanji1.xml', "常用漢字1.docx")
dump_all('kanji2.xml', "常用漢字2.docx")
dump_all('kanji3.xml', "常用漢字3.docx")
dump_all('kanji4.xml', "常用漢字4.docx")
dump_all('kanji5.xml', "常用漢字5.docx")
dump_all('kanji6.xml', "常用漢字6.docx")
dump_all('kanji7.xml', "常用漢字7.docx")
dump_all('kanji8.xml', "常用漢字8.docx")
dump_all('kanji9.xml', "常用漢字9.docx")
dump_all('kanji10.xml', "常用漢字10.docx")
dump_all('kanji11.xml', "常用漢字11.docx")
dump_all('kanji12.xml', "常用漢字12.docx")
dump_all('kanji13.xml', "常用漢字13.docx")
dump_all('kanji14.xml', "常用漢字14.docx")
dump_all('kanji15.xml', "常用漢字15.docx")
dump_all('kanji16.xml', "常用漢字16.docx")
dump_all('kanji17.xml', "常用漢字17.docx")
dump_all('kanji18.xml', "常用漢字18.docx")
dump_all('kanji19.xml', "常用漢字19.docx")
dump_all('kanji20.xml', "常用漢字20.docx")
dump_all('kanji21.xml', "常用漢字21.docx")
dump_all('kanji22.xml', "常用漢字22.docx")
dump_all('kanji23.xml', "常用漢字23.docx")
print(count)
"""

kanji_list = []

def count_kanji(doc, kanji):
	global kanji_list
	paragragh = doc.add_paragraph()
	run = paragragh.add_run(kanji.attrib['id'])
	print(kanji.attrib['id'])
	kanji_list.append(kanji.attrib['id'])
	
def count_all(xml_name):
	doc = Document()
	f = open(xml_name, 'rb')
	file_context=f.read() 
	f.close()
	
	tree = ET.ElementTree(file=xml_name)
	root = tree.getroot()
	#print(root.tag, root.attrib)
	for child in root:
		count_kanji(doc, child)

count_all('kanji1.xml')
count_all('kanji2.xml')
count_all('kanji3.xml')
count_all('kanji4.xml')
count_all('kanji5.xml')
count_all('kanji6.xml')
count_all('kanji7.xml')
count_all('kanji8.xml')
count_all('kanji9.xml')
count_all('kanji10.xml')
count_all('kanji11.xml')
count_all('kanji12.xml')
count_all('kanji13.xml')
count_all('kanji14.xml')
count_all('kanji15.xml')
count_all('kanji16.xml')
count_all('kanji17.xml')
count_all('kanji18.xml')
count_all('kanji19.xml')
count_all('kanji20.xml')
count_all('kanji21.xml')
count_all('kanji22.xml')
count_all('kanji23.xml')

print(len(kanji_list))

